import streamlit as st
import requests
from openai import OpenAI
from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
import io

# Initialize OpenAI client
client = OpenAI(api_key="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")

# Define the sections and their corresponding placeholders
SECTIONS = {
    "GenAI": "[GENAI UPDATES]",
    "AI": "[AI INSIGHTS UPDATES]",
    "IA": "[INTELLIGENT AUTOMATION UPDATES]"
}

def generate_search_query(question, client):
    try:
        # Get current month and year
        current_date = datetime.now()
        current_month = current_date.strftime("%B")
        current_year = current_date.year

        prompt = f"Generate a concise search query for Google News based on the following question in {current_month} {current_year} in europe and germany. Do not use any quotation marks in the query:\n\nQuestion: {question}\n\nSearch Query:"
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        query = response.choices[0].message.content.strip()
        # Additional check to remove any quotation marks that might have been generated
        query = query.replace('"', '').replace("'", "")
        return query
    except Exception as e:
        st.error(f"Error generating search query: {e}")
        return None
    
def fetch_news(query, num_articles=5):
    url = f"https://news.google.com/rss/search?q={query}&hl=en-US&gl=US&ceid=US:en"
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'xml')
        items = soup.find_all('item')
        articles = []
        for item in items[:num_articles]:
            title = item.title.text if item.title else "No title"
            link = item.link.text if item.link else "No link"
            articles.append({"title": title, "link": link})
        return articles
    except requests.RequestException as e:
        st.error(f"Error fetching news for query '{query}': {e}")
        return []

def summarize_article(article, client):
    try:
        prompt = f"Summarize this news article in a very short paragraph:\n\nTitle: {article['title']}\nURL: {article['link']}"
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        summary = response.choices[0].message.content.strip()
        return summary
    except Exception as e:
        st.error(f"Error summarizing article: {e}")
        return f"Unable to summarize this article due to an error: {str(e)}"

def generate_section_content(question, articles, client):
    content = ""
    for i, article in enumerate(articles, 1):
        summary = summarize_article(article, client)
        content += f"\n{article['title']}\n"
        content += f"Summary: {summary}\n"
        content += f"Read more: {article['link']}\n\n"
    return content

def process_questions(questions_and_sections):
    generated_content = {section: "" for section in SECTIONS.keys()}
    for question, section, num_articles in questions_and_sections:
        query = generate_search_query(question, client)
        if query:
            articles = fetch_news(query, num_articles)
            if articles:
                content = generate_section_content(question, articles, client)
                generated_content[section] += content + "\n\n"
            else:
                st.warning(f"No articles found for question '{question}'. Skipping this question.")
        else:
            st.warning(f"Failed to generate search query for question '{question}'. Skipping this question.")
    
    # Debug print
    st.write("Debug - Generated Content:")
    for section, content in generated_content.items():
        st.write(f"{section}: {'Content generated' if content else 'No content'}")
    
    return generated_content
    
def replace_text_in_run(run, replacements, log):
    replacement_made = False
    for key, value in replacements.items():
        if key in run.text:
            original_text = run.text
            run.text = original_text.replace(key, value)
            replacement_made = True
            log.append(f"Replaced '{key}' in run: '{original_text}' -> '{run.text}'")
    return replacement_made

def process_paragraph(paragraph, replacements, log):
    replacement_made = False
    for run in paragraph.runs:
        if replace_text_in_run(run, replacements, log):
            replacement_made = True
    return replacement_made

def process_table(table, replacements, log, table_index):
    replacement_made = False
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            # Log cell content
            log.append(f"Table {table_index}, Row {row_index}, Cell {cell_index} content: '{cell.text}'")
            
            # Process paragraphs in the cell
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                if process_paragraph(paragraph, replacements, log):
                    replacement_made = True
                    log.append(f"Replacement made in Table {table_index}, Row {row_index}, Cell {cell_index}, Paragraph {paragraph_index}")
            
            # Process nested tables
            for nested_table_index, nested_table in enumerate(cell.tables):
                if process_table(nested_table, replacements, log, f"{table_index}.{nested_table_index}"):
                    replacement_made = True
    return replacement_made

def replace_placeholders(doc, replacements):
    log = []
    replacement_made = False
    
    # Process paragraphs in the main document
    for i, paragraph in enumerate(doc.paragraphs):
        if process_paragraph(paragraph, replacements, log):
            replacement_made = True
            log.append(f"Replacement made in paragraph {i}")
    
    # Process tables in the main document
    for i, table in enumerate(doc.tables):
        if process_table(table, replacements, log, i):
            replacement_made = True
    
    return replacement_made, log

def main():
    st.set_page_config(page_title="Question-Based News Digest Generator", page_icon="📰", layout="wide")
    st.title("📰 Multi-Topic Question-Based News Digest Generator")

    st.sidebar.header("Configuration")
    
    uploaded_file = st.sidebar.file_uploader("Upload Word Template", type="docx")
    
    if uploaded_file is not None:
        template_doc = uploaded_file.getvalue()
        st.success("Template uploaded successfully!")
    else:
        st.warning("Please upload a Word template file to continue.")
        return

    num_questions = st.sidebar.number_input("Number of Questions", min_value=1, max_value=5, value=2)

    questions_and_sections = []

    # Define questions and sections in sidebar
    for i in range(num_questions):
        st.sidebar.subheader(f"Question {i+1}")
        question = st.sidebar.text_area(f"Enter question {i+1}", value=f"What are the latest developments in AI?", key=f"question_{i}")
        section = st.sidebar.selectbox(f"Section for question {i+1}", options=list(SECTIONS.keys()), key=f"section_{i}")
        num_articles = st.sidebar.number_input(f"Number of articles for question {i+1}", min_value=1, max_value=10, value=3)
        questions_and_sections.append((question, section, num_articles))
        
    if st.sidebar.button("Generate Newsletter"):
        with st.spinner("Generating newsletter..."):
            generated_content = process_questions(questions_and_sections)

        if not any(generated_content.values()):
            st.error("No articles found for any questions. Unable to generate newsletter.")
            return

        # Prepare replacements dictionary
        replacements = {SECTIONS[section]: content.strip() for section, content in generated_content.items() if content.strip()}
        replacements['[DATE]'] = datetime.now().strftime('%B %Y')

        # Debug print
        st.write("Debug - Replacements:")
        for placeholder, content in replacements.items():
            st.write(f"{placeholder}: {'Content available' if content else 'No content'}")

        # Generate the newsletter in Word format
        doc = Document(io.BytesIO(template_doc))
        replacements_made, log = replace_placeholders(doc, replacements)
        
        if not replacements_made:
            st.warning("No placeholders were replaced in the document. Please check your template and placeholders.")
        else:
            st.success("Placeholders replaced successfully.")
        
        # Display detailed log
        st.write("Replacement Log:")
        for entry in log:
            st.write(entry)
        
        # Save the document to a BytesIO object
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        st.success("Newsletter generated successfully!")

        # Offer the document for download
        st.download_button(
            label="Download Newsletter as Word Document",
            data=doc_io,
            file_name="question_based_newsletter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
